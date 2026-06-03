"""
Assemble the chapter markdown files into two finished Word documents
(Report 1 and Report 2), preserving the inline citations and per-chapter
reference lists exactly as written so the data and its sources stay together.

Each report gets: a title page, a table of contents, then every chapter in
order. Inline markdown (**bold**, *italic*, `code`), tables, headings, lists,
blockquotes, and horizontal rules are rendered.

Usage:
    python scripts/build_reports_docx.py
Outputs:
    report/Report-1-Iron-Steel-Copper-Automotive.docx
    report/Report-2-Textiles-Pharma-Petrochem-Sugar-Plastics-Cement.docx
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).parent.parent
CH   = BASE / "report" / "chapters"

DARK_BLUE = RGBColor(0x00, 0x2B, 0x5B)
MID_BLUE  = RGBColor(0x00, 0x5C, 0xA8)
GREY      = RGBColor(0x55, 0x55, 0x55)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

REPORTS = {
    "Report-1-Iron-Steel-Copper-Automotive.docx": {
        "title": "Industrial Diagnostic Study — Report 1",
        "subtitle": "Iron & Steel · Copper & Allied Metals · Automotive Value Chains",
        "glob": "report1-*.md",
    },
    "Report-2-Textiles-Pharma-Petrochem-Sugar-Plastics-Cement.docx": {
        "title": "Industrial Diagnostic Study — Report 2",
        "subtitle": ("Textiles & Garments · Pharmaceuticals · Petrochemicals & "
                     "Fertilizers · Sugar & Confectionery · Plastics & Packaging · "
                     "Cement & Building Materials"),
        "glob": "report2-*.md",
    },
}

CLIENT = "Ministry of Trade, Industry and Cooperatives (MTIC), Uganda"
PROG   = ("Uganda–UNIDO Programme for Country Partnership (PCP) for Inclusive "
          "and Sustainable Industrial Development (ISID)")


# ── inline + table helpers ──────────────────────────────────────────────────

def add_inline(paragraph, text):
    """Render **bold**, *italic*, `code` into runs on an existing paragraph."""
    for part in re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1]); r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
        else:
            paragraph.add_run(part)


def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_sep(line):
    return bool(re.match(r'^\|[\s:|-]+\|?$', line.strip())) and "-" in line


def add_table(doc, rows):
    rows = [r for r in rows if not is_sep(r)]
    if not rows:
        return
    cells = [split_row(r) for r in rows]
    ncol = max(len(c) for c in cells)
    t = doc.add_table(rows=len(cells), cols=ncol)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ri, row in enumerate(cells):
        for ci in range(ncol):
            cell = t.cell(ri, ci)
            val = row[ci] if ci < len(row) else ""
            val = val.replace("<br>", "\n")
            p = cell.paragraphs[0]
            # render first line inline; extra lines as plain breaks
            segs = val.split("\n")
            add_inline(p, segs[0])
            for extra in segs[1:]:
                pp = cell.add_paragraph(); add_inline(pp, extra)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
                if ri == 0:
                    run.font.bold = True; run.font.color.rgb = WHITE
            shade(cell, "002B5B" if ri == 0 else ("EEF3F9" if ri % 2 else "FFFFFF"))
    doc.add_paragraph()


# ── markdown → docx for one chapter ─────────────────────────────────────────

def render_markdown(doc, md):
    lines = md.split("\n")
    i = 0
    tbuf = []
    while i < len(lines):
        line = lines[i].rstrip("\n")
        if line.lstrip().startswith("|"):
            tbuf.append(line); i += 1; continue
        if tbuf:
            add_table(doc, tbuf); tbuf = []
        s = line.strip()
        if not s:
            i += 1; continue
        if s == "---":
            i += 1; continue
        if s.startswith("#### "):
            p = doc.add_heading(level=4); add_inline(p, s[5:])
        elif s.startswith("### "):
            p = doc.add_heading(level=3); add_inline(p, s[4:])
        elif s.startswith("## "):
            p = doc.add_heading(level=2); add_inline(p, s[3:])
        elif s.startswith("# "):
            p = doc.add_heading(level=1); add_inline(p, s[2:])
        elif s.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote"); add_inline(p, s[2:])
        elif re.match(r'^\d+\.\s', s):
            p = doc.add_paragraph(style="List Number"); add_inline(p, re.sub(r'^\d+\.\s', '', s))
        elif s.startswith("- ") or s.startswith("* "):
            p = doc.add_paragraph(style="List Bullet"); add_inline(p, s[2:])
        else:
            p = doc.add_paragraph(); add_inline(p, s)
        i += 1
    if tbuf:
        add_table(doc, tbuf)


# ── document scaffolding ────────────────────────────────────────────────────

def style_base(doc):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"; st.font.size = Pt(11)
    for sec in doc.sections:
        sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(2.2)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)


def title_page(doc, title, subtitle):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = DARK_BLUE
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle); r.font.size = Pt(13); r.font.color.rgb = MID_BLUE
    doc.add_paragraph()
    for label in (CLIENT, PROG):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label); r.font.size = Pt(11); r.font.color.rgb = GREY
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Draft for validation"); r.italic = True; r.font.size = Pt(11)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Sources and citations are provided inline and in per-chapter "
                  "reference lists for verification."); r.font.size = Pt(9.5); r.font.color.rgb = GREY
    doc.add_page_break()


def toc_field(doc):
    """Insert an auto-updating Word table of contents (updates on open/F9)."""
    h = doc.add_heading(level=1); add_inline(h, "Table of Contents")
    p = doc.add_paragraph()
    run = p.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin"); run._r.append(fb)
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = r'TOC \o "1-2" \h \z \u'; run._r.append(it)
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate"); run._r.append(fs)
    t = OxmlElement("w:t"); t.text = "Right-click and choose “Update Field” to build the contents."
    run._r.append(t)
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end"); run._r.append(fe)
    doc.add_page_break()


def build(out_name, cfg):
    doc = Document()
    style_base(doc)
    title_page(doc, cfg["title"], cfg["subtitle"])
    toc_field(doc)
    files = sorted(CH.glob(cfg["glob"]))
    for idx, f in enumerate(files):
        render_markdown(doc, f.read_text(encoding="utf-8"))
        if idx < len(files) - 1:
            doc.add_page_break()
    out = BASE / "report" / out_name
    doc.save(out)
    return out, len(files)


if __name__ == "__main__":
    for name, cfg in REPORTS.items():
        out, n = build(name, cfg)
        print(f"Built {out.name}  ({n} chapters)")
    print("Done.")
