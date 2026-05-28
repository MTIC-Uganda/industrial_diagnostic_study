"""
Convert inception-report-1.md and inception-report-2.md to .docx
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK_BLUE = RGBColor(0x00, 0x2B, 0x5B)
MID_BLUE  = RGBColor(0x00, 0x5C, 0xA8)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG  = RGBColor(0xD6, 0xE8, 0xF7)


def set_heading_style(para, level, text):
    para.clear()
    run = para.add_run(text)
    run.bold = True
    if level == 0:
        run.font.size = Pt(22)
        run.font.color.rgb = DARK_BLUE
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = DARK_BLUE
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = MID_BLUE
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after  = Pt(4)


def apply_inline(run_text, para):
    """Parse **bold** and *italic* inline and add runs to para."""
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', run_text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('*') and part.endswith('*'):
            r = para.add_run(part[1:-1])
            r.italic = True
        elif part.startswith('`') and part.endswith('`'):
            r = para.add_run(part[1:-1])
            r.font.name = 'Courier New'
            r.font.size = Pt(9)
        else:
            para.add_run(part)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_table(doc, headers, rows):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        shade_cell(cell, '002B5B')
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)

    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        fill = 'D6E8F7' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            shade_cell(cell, fill)
            p = cell.paragraphs[0]
            apply_inline(val, p)
            for r in p.runs:
                r.font.size = Pt(10)

    doc.add_paragraph()


def parse_table_line(line):
    parts = [c.strip() for c in line.strip().strip('|').split('|')]
    return parts


def is_separator(line):
    return bool(re.match(r'^\|[-| :]+\|$', line.strip()))


def convert(md_path, docx_path):
    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default paragraph font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    i = 0
    table_buf = []
    in_table  = False
    in_code   = False
    code_buf  = []

    while i < len(lines):
        raw  = lines[i]
        line = raw.rstrip('\n')

        # ── Code block (Gantt chart) ────────────────────────────
        if line.strip().startswith('```'):
            if not in_code:
                in_code  = True
                code_buf = []
            else:
                in_code = False
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(4)
                r = p.add_run('\n'.join(code_buf))
                r.font.name = 'Courier New'
                r.font.size = Pt(7.5)
                doc.add_paragraph()
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ── Table ───────────────────────────────────────────────
        if line.startswith('|'):
            table_buf.append(line)
            i += 1
            continue
        else:
            if table_buf:
                rows_raw = [l for l in table_buf if not is_separator(l)]
                if rows_raw:
                    headers = parse_table_line(rows_raw[0])
                    data    = [parse_table_line(r) for r in rows_raw[1:]]
                    add_table(doc, headers, data)
                table_buf = []

        # ── Horizontal rule ─────────────────────────────────────
        if line.strip() == '---':
            doc.add_paragraph()
            i += 1
            continue

        # ── Headings ────────────────────────────────────────────
        h1 = re.match(r'^# (.+)$', line)
        h2 = re.match(r'^## (.+)$', line)
        h3 = re.match(r'^### (.+)$', line)
        h4 = re.match(r'^#### (.+)$', line)

        if h1:
            p = doc.add_heading('', level=1)
            set_heading_style(p, 0, h1.group(1))
        elif h2:
            p = doc.add_heading('', level=2)
            set_heading_style(p, 1, h2.group(1))
        elif h3:
            p = doc.add_heading('', level=3)
            set_heading_style(p, 2, h3.group(1))
        elif h4:
            p = doc.add_heading('', level=4)
            set_heading_style(p, 2, h4.group(1))

        # ── Metadata lines (bold key: value) ────────────────────
        elif re.match(r'^\*\*[^*]+:\*\*', line):
            p = doc.add_paragraph()
            apply_inline(line, p)
            p.paragraph_format.space_after = Pt(2)

        # ── Blockquote ──────────────────────────────────────────
        elif line.startswith('> '):
            p = doc.add_paragraph(style='Quote')
            p.paragraph_format.left_indent = Inches(0.4)
            text = line[2:].strip().strip('*')
            r = p.add_run(text)
            r.italic = True
            r.font.size = Pt(11)

        # ── Numbered list ───────────────────────────────────────
        elif re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(style='List Number')
            apply_inline(re.sub(r'^\d+\. ', '', line), p)
            p.paragraph_format.space_after = Pt(2)

        # ── Bullet list ─────────────────────────────────────────
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            apply_inline(line[2:], p)
            p.paragraph_format.space_after = Pt(2)

        # ── Part headings (bold paragraph) ──────────────────────
        elif line.startswith('### Part') or (line.startswith('**') and line.endswith('**')):
            p = doc.add_paragraph()
            apply_inline(line, p)
            p.paragraph_format.space_before = Pt(8)

        # ── Empty line ──────────────────────────────────────────
        elif line.strip() == '':
            doc.add_paragraph()

        # ── Normal paragraph ────────────────────────────────────
        else:
            p = doc.add_paragraph()
            apply_inline(line, p)
            p.paragraph_format.space_after = Pt(4)

        i += 1

    # Flush any remaining table
    if table_buf:
        rows_raw = [l for l in table_buf if not is_separator(l)]
        if rows_raw:
            headers = parse_table_line(rows_raw[0])
            data    = [parse_table_line(r) for r in rows_raw[1:]]
            add_table(doc, headers, data)

    doc.save(docx_path)
    print(f'Saved: {docx_path}')


if __name__ == '__main__':
    base = r'd:\graphics\Jero\MTIC\industrial_diagnostic_study\report'
    convert(f'{base}\\inception-report-1.md', f'{base}\\inception-report-1.docx')
    convert(f'{base}\\inception-report-2.md', f'{base}\\inception-report-2.docx')
    print('Both Word documents created.')
