#!/usr/bin/env python3
"""
MTIC Ingestion Agent — Phase 2

Triggered by GitHub Actions whenever Jerome pushes a file to data/uploads/.

Three-method protocol (per docs/Ingestion_Specification.md):
  1. upload_mining  — mine unstructured docs Jerome uploaded
  2. trusted_source — pull authoritative figures from trusted_sources.json registry
  3. web_gap_fill   — propose candidates for remaining mandatory gaps; opens a PR

Output: seeds approved datapoints to PocketBase collection `diagnostic_datapoints`.
"""

import json, os, sys, time, datetime, hashlib, textwrap, urllib.request, urllib.error
from pathlib import Path
import sys as _sys; _sys.path.insert(0, str(Path(__file__).parent))
from _status import update_status

ROOT        = Path(__file__).resolve().parent.parent
SCHEMA_FILE = ROOT / 'data' / 'schema' / 'diagnostic_schema.json'
SOURCES_FILE= ROOT / 'data' / 'schema' / 'trusted_sources.json'
UPLOADS_DIR = ROOT / 'data' / 'uploads'
INGESTED_DIR= ROOT / 'data' / 'ingested'
PROPOSED_DIR= ROOT / 'data' / 'proposed'

PB_URL   = os.environ.get('PB_URL', 'http://89.167.121.193:8090').rstrip('/')
PB_EMAIL = os.environ.get('PB_ADMIN_EMAIL', '')
PB_PASS  = os.environ.get('PB_ADMIN_PASSWORD', '')
ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY', '')

RUN_ID   = datetime.datetime.utcnow().strftime('%Y%m%dT%H%M%SZ')
TODAY    = datetime.date.today().isoformat()

INGESTED_DIR.mkdir(parents=True, exist_ok=True)
PROPOSED_DIR.mkdir(parents=True, exist_ok=True)

# ── Load schema ───────────────────────────────────────────────────────────────

schema     = json.loads(SCHEMA_FILE.read_text('utf-8'))
sources_db = json.loads(SOURCES_FILE.read_text('utf-8'))

FIELDS     = {f['field_id']: f for f in schema['fields']}
VC_IDS     = {vc['name']: vc['id'] for vc in schema['value_chains']}
VC_NAMES   = {vc['id']: vc['name'] for vc in schema['value_chains']}
MANDATORY  = [f['field_id'] for f in schema['fields'] if f['required_level'] == 'mandatory']

# ── PocketBase auth ───────────────────────────────────────────────────────────

TOKEN = None

def _req(method, path, payload=None, token=None):
    url  = f'{PB_URL}/{path.lstrip("/")}'
    data = json.dumps(payload).encode() if payload is not None else None
    hdrs = {'Content-Type': 'application/json'}
    if token:
        hdrs['Authorization'] = token
    r = urllib.request.Request(url, data=data, method=method, headers=hdrs)
    try:
        with urllib.request.urlopen(r) as resp:
            body = resp.read()
            return json.loads(body) if body else {}
    except urllib.error.HTTPError as e:
        body = e.read().decode()
        print(f'  HTTP {e.code} {method} {path}: {body[:200]}')
        raise

def pb(method, path, payload=None):
    return _req(method, path, payload, token=TOKEN)

def pb_auth():
    global TOKEN
    if not PB_EMAIL or not PB_PASS:
        print('WARN: PB_ADMIN_EMAIL/PB_ADMIN_PASSWORD not set — skipping PocketBase seed.')
        return False
    for endpoint in ['/api/admins/auth-with-password',
                     '/api/collections/_superusers/auth-with-password']:
        try:
            auth = _req('POST', endpoint, {'identity': PB_EMAIL, 'password': PB_PASS})
            TOKEN = auth['token']
            print(f'  PocketBase: authenticated as {PB_EMAIL}')
            return True
        except urllib.error.HTTPError:
            continue
    print('WARN: PocketBase auth failed — will write JSON output only.')
    return False

def pb_seed(datapoints):
    """Upsert each datapoint by (field_id, value_chain_id) key."""
    seeded = 0
    for dp in datapoints:
        fid = dp['field_id']
        vc  = dp['value_chain_id']
        existing = pb('GET',
            f'/api/collections/diagnostic_datapoints/records'
            f'?filter=(field_id="{fid}"&&value_chain_id="{vc}")&perPage=1'
        ).get('items', [])
        if existing:
            pb('PATCH', f'/api/collections/diagnostic_datapoints/records/{existing[0]["id"]}', dp)
        else:
            pb('POST', '/api/collections/diagnostic_datapoints/records', dp)
        seeded += 1
    print(f'  PocketBase: seeded {seeded} datapoints')

# ── Claude API call ───────────────────────────────────────────────────────────

def claude(system_prompt, user_prompt, model=None):
    """
    Call the Claude CLI installed on the Hetzner host (the standing setup; there
    is no Anthropic API key). Runs headless: prompt in on stdin, text answer out.
    The CLI carries the Max-plan auth under HOME=/root, the same way
    /opt/whatsapp-agent/code_dispatch.py drives it. CLAUDE_BIN lets a different
    host override the path.
    """
    import subprocess
    claude_bin = os.environ.get('CLAUDE_BIN', 'claude')
    full = f'{system_prompt}\n\n{user_prompt}'
    env = dict(os.environ)
    env.setdefault('HOME', '/root')
    cmd = [claude_bin, '-p', '--output-format', 'text']
    if model:
        cmd += ['--model', model]
    try:
        r = subprocess.run(cmd, input=full, capture_output=True, text=True,
                           env=env, timeout=600)
    except FileNotFoundError:
        sys.exit('ERROR: Claude CLI not found. Set CLAUDE_BIN or install the CLI '
                 '(this agent runs on Hetzner, not in CI).')
    if r.returncode != 0:
        raise RuntimeError(f'Claude CLI failed (exit {r.returncode}): {r.stderr[:500]}')
    return r.stdout.strip()

# ── System prompt shared by all mining calls ──────────────────────────────────

SCHEMA_SUMMARY = json.dumps({
    'fields': [
        {'field_id': f['field_id'], 'label': f['label'], 'domain': f['domain'],
         'data_type': f['data_type'], 'required_level': f['required_level'],
         'validation_rule': f['validation_rule'],
         'synthesis_guidance': f['synthesis_guidance']}
        for f in schema['fields']
    ],
    'value_chains': schema['value_chains'],
    'controlled_vocabularies': schema['controlled_vocabularies'],
}, ensure_ascii=False)

SYSTEM_PROMPT = textwrap.dedent(f"""
You are the MTIC Ingestion Agent for Uganda's Industrial Diagnostic Study.

Your job is to extract factual data from source documents and map every fact to the
correct field in the diagnostic schema. You collect data using upload_mining:
you read the document and extract values that are explicitly stated or can be directly
computed from stated figures. You never infer or fabricate.

CONTRACT:
- For every field you can populate, produce one JSON object matching the record envelope.
- If a mandatory field has no data in the document, set value="not_available" and
  explain in collection_notes.
- Set ingestion_method="upload_mining" and approval_status="not_required" for all records.
- Set confidence honestly: high = directly stated with source, medium = computed/triangulated,
  low = approximate or indirect, estimated = desk estimate, not_available = no data found.
- Record source_document_ref as: "<filename>, <page or section>".
- evidence_required fields MUST have a raw_source citation.
- Never invent a figure. Unknown = not_available + reason in collection_notes.

SCHEMA (66 fields across 12 domains):
{SCHEMA_SUMMARY}

RECORD ENVELOPE (output format for each datapoint):
{{
  "field_id": "D1.01",
  "value_chain_id": "VC01",
  "value": "<extracted value>",
  "raw_source": "<document title, page/section, exact quote if numeric>",
  "source_type": "<from controlled vocabulary>",
  "reporting_year": "<year or period>",
  "confidence": "<high|medium|low|estimated|not_available>",
  "disaggregation_values": "<any splits captured, or null>",
  "collection_notes": "<caveats, assumptions, data limits>",
  "collected_by": "ingestion_agent/{RUN_ID}",
  "triangulation_refs": null,
  "ingestion_method": "upload_mining",
  "trusted_source_id": null,
  "source_document_ref": "<filename, page/section>",
  "source_url": null,
  "approval_status": "not_required",
  "approved_by": null,
  "approved_at": null,
  "approval_pr": null
}}

OUTPUT: Return ONLY a JSON array of record envelope objects. No prose, no markdown fences.
""").strip()

# ── Document text extraction ──────────────────────────────────────────────────

def extract_text(filepath: Path) -> str:
    """
    Extract readable text from a source document.

    Handles the formats Jerome actually uploads — PDF, DOCX, XLSX — plus plain
    text/markdown/CSV/HTML. The old code used filepath.read_text('utf-8') for
    everything, which turned a PDF into mojibake and silently broke mining. We
    extract per type and raise a clear error on failure rather than feed garbage
    to Claude.
    """
    suffix = filepath.suffix.lower()
    if suffix == '.pdf':
        try:
            import pdfplumber
            with pdfplumber.open(filepath) as pdf:
                return '\n'.join((page.extract_text() or '') for page in pdf.pages)
        except Exception:
            from pypdf import PdfReader            # fallback extractor
            reader = PdfReader(str(filepath))
            return '\n'.join((page.extract_text() or '') for page in reader.pages)
    if suffix == '.docx':
        try:
            from docx import Document
            doc = Document(str(filepath))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:               # tables hold most of the data
                for row in table.rows:
                    parts.append('\t'.join(cell.text for cell in row.cells))
            return '\n'.join(parts)
        except Exception:
            # Some DOCX files have malformed internal paths that break python-docx.
            # Fall back to pulling text straight from word/document.xml in the zip.
            import zipfile, re
            with zipfile.ZipFile(str(filepath)) as z:
                xml = z.read('word/document.xml').decode('utf-8', errors='replace')
            xml = re.sub(r'</w:p>', '\n', xml)     # paragraph breaks
            return re.sub(r'<[^>]+>', '', xml)     # strip all tags
    if suffix in ('.xlsx', '.xlsm'):
        from openpyxl import load_workbook
        wb = load_workbook(str(filepath), read_only=True, data_only=True)
        lines = []
        for ws in wb.worksheets:
            lines.append(f'# Sheet: {ws.title}')
            for row in ws.iter_rows(values_only=True):
                cells = ['' if c is None else str(c) for c in row]
                if any(cells):
                    lines.append('\t'.join(cells))
        return '\n'.join(lines)
    # txt, md, csv, html, and anything else: read as text
    return filepath.read_text(encoding='utf-8', errors='replace')

# ── Method 1: Upload mining ───────────────────────────────────────────────────

def read_task_sidecar(filepath: Path) -> str:
    """
    Read a per-document instruction sidecar `<document>.task.md` if present.

    Intent travels with the document instead of living in a PR body the agent
    never reads (ADR-010). Returns '' when no sidecar exists.
    """
    sidecar = filepath.with_suffix(filepath.suffix + '.task.md')
    if sidecar.exists():
        return sidecar.read_text(encoding='utf-8', errors='replace').strip()
    return ''

def mine_upload(filepath: Path, value_chain_id: str) -> list[dict]:
    """Send file content to Claude and get back structured datapoints."""
    vc_name = VC_NAMES.get(value_chain_id, value_chain_id)
    content = extract_text(filepath)

    # Truncate very large files — keep first 80k chars (Claude context limit buffer)
    if len(content) > 80_000:
        content = content[:80_000] + '\n\n[... document truncated at 80,000 chars ...]'

    task = read_task_sidecar(filepath)
    task_block = f'\n    OPERATOR INSTRUCTION (from {filepath.name}.task.md):\n    {task}\n' if task else ''

    user_prompt = textwrap.dedent(f"""
    VALUE CHAIN: {vc_name} ({value_chain_id})
    SOURCE FILE: {filepath.name}
    {task_block}
    Extract all diagnostic data you can find for the {vc_name} value chain from this document.
    Map each fact to the appropriate field_id from the schema.
    Return a JSON array of record envelope objects.

    DOCUMENT CONTENT:
    {content}
    """).strip()

    print(f'  Mining {filepath.name} for {vc_name}...')
    raw = claude(SYSTEM_PROMPT, user_prompt)

    # Strip any accidental markdown fences
    raw = raw.strip()
    if raw.startswith('```'):
        raw = raw.split('\n', 1)[1].rsplit('```', 1)[0]

    try:
        datapoints = json.loads(raw)
    except json.JSONDecodeError as e:
        print(f'  WARN: Claude returned invalid JSON for {filepath.name}: {e}')
        datapoints = []

    # Stamp collected_by
    for dp in datapoints:
        dp['collected_by'] = f'ingestion_agent/{RUN_ID}'

    return datapoints

# ── Method 2: Trusted sources ─────────────────────────────────────────────────

def pull_trusted_sources(value_chain_id: str) -> list[dict]:
    """
    Pull figures from sources registered in trusted_sources.json.
    Currently implements TradeMap via the ITC public CSV endpoint.
    Other sources (UBOS, Comtrade) are stubbed — extend as APIs become available.
    """
    datapoints = []
    for source in sources_db.get('sources', []):
        sid = source['source_id']
        authoritative_for = source.get('authoritative_for', [])
        if not authoritative_for:
            continue

        if sid == 'trade_map':
            # TradeMap figures are already in data/dashboard/ CSVs (scraped separately).
            # Read from the chain summary CSV as the trusted source.
            csv_path = ROOT / 'data' / 'dashboard' / 'chain_summary.csv'
            if not csv_path.exists():
                continue
            import csv
            vc_name = VC_NAMES.get(value_chain_id, '')
            with open(csv_path, newline='', encoding='utf-8') as f:
                for row in csv.DictReader(f):
                    if row.get('chain', '').strip() == vc_name:
                        for field_id in authoritative_for:
                            # Map chain_summary columns to field_ids
                            col_map = {
                                'D1.02': 'key_import_2024',
                                'D1.03': 'key_export_2024',
                            }
                            col = col_map.get(field_id)
                            if col and row.get(col):
                                datapoints.append({
                                    'field_id': field_id,
                                    'value_chain_id': value_chain_id,
                                    'value': row[col],
                                    'raw_source': 'ITC TradeMap — imported via chain_summary.csv',
                                    'source_type': 'trusted_api',
                                    'reporting_year': '2024',
                                    'confidence': source.get('default_confidence', 'high'),
                                    'disaggregation_values': None,
                                    'collection_notes': None,
                                    'collected_by': f'ingestion_agent/{RUN_ID}',
                                    'triangulation_refs': None,
                                    'ingestion_method': 'trusted_source',
                                    'trusted_source_id': sid,
                                    'source_document_ref': None,
                                    'source_url': source.get('base_url'),
                                    'approval_status': 'not_required',
                                    'approved_by': None,
                                    'approved_at': None,
                                    'approval_pr': None,
                                })
        # Add more source integrations here as needed (UBOS, Comtrade, etc.)

    return datapoints

# ── Method 3: Web gap-fill ────────────────────────────────────────────────────

def detect_gaps(collected: list[dict], value_chain_id: str) -> list[str]:
    """Return mandatory field_ids not yet populated for this chain."""
    populated = {dp['field_id'] for dp in collected
                 if dp.get('value') and dp.get('value') != 'not_available'}
    return [fid for fid in MANDATORY if fid not in populated]

def web_gap_fill(gaps: list[str], value_chain_id: str) -> list[dict]:
    """
    For remaining mandatory gaps, ask Claude to propose candidate values
    from its training knowledge. These get approval_status=pending_approval
    and are written to data/proposed/ — never seeded to PocketBase directly.
    A PR must be opened and merged before they enter the dataset.
    """
    if not gaps:
        return []

    vc_name = VC_NAMES.get(value_chain_id, value_chain_id)
    gap_fields = [
        {'field_id': fid, 'label': FIELDS[fid]['label'],
         'data_type': FIELDS[fid]['data_type'],
         'synthesis_guidance': FIELDS[fid].get('synthesis_guidance', '')}
        for fid in gaps if fid in FIELDS
    ]

    user_prompt = textwrap.dedent(f"""
    VALUE CHAIN: {vc_name} ({value_chain_id})

    The following mandatory fields could not be populated from uploaded documents or
    trusted sources. Propose candidate values based on your knowledge of Uganda's
    industrial sector. Be conservative — prefer not_available over a guess you are
    not confident in.

    GAPS TO FILL:
    {json.dumps(gap_fields, indent=2)}

    For each gap, return a record envelope with:
    - ingestion_method: "web_gap_fill"
    - approval_status: "pending_approval"
    - source_url: the most relevant public URL where this figure can be verified
    - confidence: your honest assessment
    - collection_notes: explain your reasoning and any caveats

    Return ONLY a JSON array. These are PROPOSALS — a human must review and merge a PR
    before they enter the dataset.
    """).strip()

    print(f'  Web gap-fill: {len(gaps)} gaps for {vc_name}...')
    raw = claude(SYSTEM_PROMPT, user_prompt)
    raw = raw.strip()
    if raw.startswith('```'):
        raw = raw.split('\n', 1)[1].rsplit('```', 1)[0]

    try:
        proposals = json.loads(raw)
    except json.JSONDecodeError:
        proposals = []

    for p in proposals:
        p['collected_by'] = f'ingestion_agent/{RUN_ID}'
        p['approval_status'] = 'pending_approval'
        p['ingestion_method'] = 'web_gap_fill'

    return proposals

# ── Main ──────────────────────────────────────────────────────────────────────

def run(upload_files: list[Path] | None = None, value_chain_id: str | None = None):
    """
    Run the full three-method ingestion pipeline.

    Args:
        upload_files:    list of files to mine (default: all files in data/uploads/)
        value_chain_id:  restrict to one chain (default: infer from filename or prompt)
    """
    print(f'\n══ MTIC Ingestion Agent  run={RUN_ID} ══')

    # Discover uploads
    if upload_files is None:
        upload_files = sorted(UPLOADS_DIR.glob('*'))
        upload_files = [f for f in upload_files if f.is_file() and not f.name.startswith('.')]

    if not upload_files:
        print('  No files in data/uploads/ — nothing to ingest.')
        print('  Tip: drop a document into data/uploads/ and push to trigger ingestion.')
        return

    # Determine which value chains to process.
    # Primary signal is the containing sector folder (data/<value-chain>/), which
    # is where Jerome and the README already file documents — far more reliable
    # than guessing from the filename (ADR-010). Filename keywords are a fallback.
    FOLDER_VC = {
        'iron-steel': 'VC01',
        'copper-allied-metals': 'VC02',
        'automotive': 'VC03',
        'textiles-garments': 'VC04',
        'pharmaceuticals': 'VC05',
        'petrochemicals-fertilizers': 'VC06',
        'sugar-confectionery': 'VC07',
        'plastics-packaging': 'VC08',
        'cement-building-materials': 'VC09',
    }

    def infer_vc(path: Path) -> str | None:
        # 1. Folder-based routing (most reliable)
        folder = path.parent.name.lower()
        if folder in FOLDER_VC:
            return FOLDER_VC[folder]
        # 2. Explicit VC id in the filename (e.g. VC07_sugarcane.pdf)
        name = path.stem.upper()
        for vc_id in VC_NAMES:
            if vc_id in name:
                return vc_id
        # 3. Chain-name keywords in the filename
        lower = path.stem.lower()
        kw_map = {
            'iron': 'VC01', 'steel': 'VC01',
            'copper': 'VC02',
            'automotive': 'VC03', 'vehicle': 'VC03',
            'textile': 'VC04', 'garment': 'VC04',
            'pharma': 'VC05',
            'petrochem': 'VC06', 'fertilizer': 'VC06',
            'sugar': 'VC07', 'confection': 'VC07',
            'plastic': 'VC08', 'packaging': 'VC08',
            'cement': 'VC09', 'building': 'VC09',
        }
        for kw, vc in kw_map.items():
            if kw in lower:
                return vc
        return None

    # Collect all datapoints across all files
    all_approved: list[dict] = []
    all_proposed: list[dict] = []

    chains_processed = set()

    for filepath in upload_files:
        vc = value_chain_id or infer_vc(filepath)
        if not vc:
            print(f'  SKIP {filepath.name} — cannot determine value chain. '
                  f'Rename to include VC01..VC09 or a chain keyword.')
            continue

        # Method 1: mine the upload. One unreadable document must not abort the
        # whole batch — log it and move on (ADR-010: visible gap, not a crash).
        try:
            mined = mine_upload(filepath, vc)
        except Exception as e:
            print(f'  SKIP {filepath.name} — could not extract/mine: {e}')
            continue
        all_approved.extend(mined)
        chains_processed.add(vc)
        print(f'    Mined {len(mined)} datapoints from {filepath.name}')

    # Method 2: trusted sources (run once per chain)
    for vc in chains_processed:
        trusted = pull_trusted_sources(vc)
        all_approved.extend(trusted)
        if trusted:
            print(f'  Trusted sources: {len(trusted)} datapoints for {VC_NAMES[vc]}')

    # Method 3: web gap-fill for remaining mandatory gaps per chain
    for vc in chains_processed:
        chain_collected = [dp for dp in all_approved if dp.get('value_chain_id') == vc]
        gaps = detect_gaps(chain_collected, vc)
        if gaps:
            print(f'  Gaps for {VC_NAMES[vc]}: {len(gaps)} mandatory fields missing '
                  f'({", ".join(gaps[:5])}{"..." if len(gaps) > 5 else ""})')
            proposals = web_gap_fill(gaps, vc)
            all_proposed.extend(proposals)

    # Write approved datapoints to JSON (source of truth before PocketBase)
    if all_approved:
        out_path = INGESTED_DIR / f'ingested_{RUN_ID}.json'
        out_path.write_text(json.dumps(all_approved, indent=2, ensure_ascii=False), 'utf-8')
        print(f'\n  Wrote {len(all_approved)} approved datapoints → {out_path.name}')

    # Write proposed (web_gap_fill) to data/proposed/ for PR review
    if all_proposed:
        prop_path = PROPOSED_DIR / f'web_gap_fill_{TODAY}.json'
        prop_path.write_text(json.dumps(all_proposed, indent=2, ensure_ascii=False), 'utf-8')
        print(f'  Wrote {len(all_proposed)} gap-fill proposals → {prop_path.name}')
        print('  ACTION: open a PR for the proposals file — your merge = approval.')

    # Seed approved datapoints to PocketBase
    if all_approved and pb_auth():
        pb_seed(all_approved)

    # Summary
    print(f'\n══ Done  approved={len(all_approved)}  proposed={len(all_proposed)} ══')
    if all_proposed:
        print('  Next: review data/proposed/ and open a PR to approve gap-fill data.')

    # Write pipeline status for each processed chain
    for vc in chains_processed:
        chain_dps = [dp for dp in all_approved if dp.get('value_chain_id') == vc]
        chain_props = [dp for dp in all_proposed if dp.get('value_chain_id') == vc]
        update_status('ingest', vc, {
            'run_id': RUN_ID,
            'datapoints_added': len(chain_dps),
            'gaps_proposed': len(chain_props),
            'source_files': [f.name for f in (upload_files or list(UPLOADS_DIR.glob('*')))
                             if f.suffix.lower() in {'.pdf', '.docx', '.xlsx', '.csv'}
                             and infer_vc(f) == vc],
        })


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description='MTIC Ingestion Agent')
    parser.add_argument('--vc', help='Value chain ID to process (e.g. VC01). '
                        'Default: infer from filename.')
    parser.add_argument('files', nargs='*', type=Path,
                        help='Specific files to ingest. Default: all in data/uploads/')
    args = parser.parse_args()

    files = [Path(f) for f in args.files] if args.files else None
    run(upload_files=files, value_chain_id=args.vc)
